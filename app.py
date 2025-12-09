# app.py
import streamlit as st
from PIL import Image
import easyocr
import numpy as np
from pdf2image import convert_from_bytes
import PyPDF2
from io import BytesIO
from docx import Document
import json
import re
import os
from dotenv import load_dotenv

# Optional OpenAI integration
try:
    from openai import OpenAI
    openai_available = True
except Exception:
    openai_available = False

load_dotenv()  # read .env locally if present

# -----------------------------
# Helpers
# -----------------------------
def extract_json_from_text(text):
    """Try to parse a JSON object inside text, robust to extra commentary."""
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r"\{.*\}", text, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                return None
    return None

def reader_instance():
    """Return a single EasyOCR reader instance (lazy init)."""
    global _reader
    try:
        _reader
    except NameError:
        _reader = easyocr.Reader(['en'], gpu=False)
    return _reader

def ocr_image(file_like):
    """OCR a file-like image (uploaded or camera). Returns extracted text."""
    try:
        img = Image.open(file_like).convert("RGB")
        arr = np.array(img)
        reader = reader_instance()
        result = reader.readtext(arr, detail=0)
        return "\n".join(result).strip()
    except Exception as e:
        return f"ERROR: {e}"

def ocr_pdf(pdf_file):
    """Try text extraction via PyPDF2, otherwise rasterize and OCR pages."""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
        if text.strip():
            return text.strip()
        # fallback: convert pages to images and OCR
        pdf_file.seek(0)
        images = convert_from_bytes(pdf_file.read())
        out = ""
        for img in images:
            out += ocr_image(img) + "\n"
        return out.strip()
    except Exception as e:
        return f"ERROR: {e}"

# -----------------------------
# AI evaluation (OpenAI) or fallback
# -----------------------------
def ai_evaluate(essay_text, score_scale=10):
    """
    Uses OpenAI to:
    - decide relevant criteria
    - score each criterion 1..score_scale
    - return corrected_essay, summary, explanations (dict)
    If OpenAI not available or fails, fall back to heuristics.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if openai_available and api_key:
        try:
            client = OpenAI(api_key=api_key)
            prompt = f"""
You are an expert writing evaluator. Given the essay below, do the following:
1) Identify 5-8 evaluation criteria that are most relevant to this essay (e.g. Clarity, Argument Strength, Grammar, Evidence, Organization, Tone).
2) For each criterion, give a **numeric score** from 1 to {score_scale} and a short (1-2 sentence) explanation.
3) Provide a corrected version of the essay (correct grammar/typos while preserving the author's meaning).
4) Provide a short summary analysis (2-4 sentences).
5) Provide sentence-by-sentence explanations (Teaching Mode).
Output ONLY a JSON object with keys:
- criteria: list of objects {{ "name":..., "score": int, "explanation": ... }}
- corrected_essay: string
- summary: string
- explanations: string

Essay:
\"\"\"{essay_text}\"\"\"
"""
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt}]
            )
            raw = response.choices[0].message.content
            data = extract_json_from_text(raw)
            if data:
                return data
            # else fallback to heuristic
        except Exception as e:
            # don't crash; proceed to fallback
            print("OpenAI call failed:", e)

    # Fallback heuristic evaluator (simple, transparent)
    return heuristic_evaluate(essay_text, score_scale)

def heuristic_evaluate(text, score_scale=10):
    """Simple heuristic evaluation if OpenAI isn't available."""
    # criteria candidates - dynamic-ish but deterministic
    candidates = [
        "Clarity of Ideas", "Organization & Flow", "Grammar & Mechanics",
        "Evidence & Support", "Vocabulary & Style", "Coherence"
    ]
    # basic heuristics (not perfect, but workable)
    words = text.split()
    word_count = max(1, len(words))
    sent_count = max(1, text.count('.') + text.count('!') + text.count('?'))
    avg_word_len = sum(len(w) for w in words) / word_count
    long_words_ratio = sum(1 for w in words if len(w) > 7) / word_count

    # produce scores 1..score_scale
    scores = {}
    # Clarity: more sentences and reasonable avg word length
    clarity = max(1, min(score_scale, int((10 * (1 - abs((avg_word_len - 5)/5))) * score_scale / 10)))
    grammar = max(1, min(score_scale, int(score_scale * (0.6 + 0.4 * min(1, 100/ (1 + text.count('  ')) ) ))))
    organization = max(1, min(score_scale, int(min(score_scale, sent_count / 2) )))
    evidence = max(1, min(score_scale, int(score_scale * min(1, long_words_ratio * 2))))
    vocab = max(1, min(score_scale, int(score_scale * min(1, long_words_ratio * 1.2))))
    coherence = max(1, min(score_scale, int((clarity + organization) / 2)))

    # assemble criteria list
    crits = [
        {"name":"Clarity of Ideas", "score": clarity, "explanation":"Based on sentence complexity and average word length."},
        {"name":"Organization & Flow", "score": organization, "explanation":"Estimated from sentence count and structure."},
        {"name":"Grammar & Mechanics", "score": grammar, "explanation":"Heuristic check; low confidence without full grammar model."},
        {"name":"Evidence & Support", "score": evidence, "explanation":"Estimated from presence of longer words/phrases (proxy)."},
        {"name":"Vocabulary & Style", "score": vocab, "explanation":"Estimated from word length and variety."},
        {"name":"Coherence", "score": coherence, "explanation":"Average of clarity and organization heuristics."}
    ]

    corrected = text  # no automatic corrections in fallback
    summary = "Automatic fallback evaluation used (OpenAI not available). Scores are heuristic estimates."
    explanations = "Sentence-by-sentence explanations are not available in fallback mode."

    return {
        "criteria": crits,
        "corrected_essay": corrected,
        "summary": summary,
        "explanations": explanations
    }

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="AI Essay Checker (Dynamic Criteria)", page_icon="📘", layout="wide")
st.markdown("<h1 style='text-align:center'>📘 AI Essay Evaluation System</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align:center'>Dynamic criteria (1–10) • Paste / Upload / Camera / PDF</p>", unsafe_allow_html=True)
st.write("---")

mode = st.radio("Choose Input Method:", ["📄 Paste Text", "📷 Upload Image", "📸 Camera Scan", "📑 Upload PDF / Scan"])
essay_text = ""

if mode == "📄 Paste Text":
    st.subheader("📝 Paste or type your essay")
    essay_text = st.text_area("Essay text:", height=300)

elif mode == "📷 Upload Image":
    st.subheader("📷 Upload image (photo of paper)")
    uploaded = st.file_uploader("Image file", type=["png","jpg","jpeg"])
    if uploaded:
        with st.spinner("Running OCR on image..."):
            essay_text = ocr_image(uploaded)
        st.subheader("📄 Extracted Text")
        st.text_area("Extracted text", value=essay_text, height=250)

elif mode == "📸 Camera Scan":
    st.subheader("📸 Camera — take a photo")
    cam = st.camera_input("Take a photo")
    if cam:
        with st.spinner("Running OCR on camera image..."):
            essay_text = ocr_image(cam)
        st.subheader("📄 Extracted Text")
        st.text_area("Extracted text", value=essay_text, height=250)

else:  # PDF
    st.subheader("📑 Upload PDF")
    uploaded_pdf = st.file_uploader("PDF file", type=["pdf"])
    if uploaded_pdf:
        with st.spinner("Extracting PDF text..."):
            essay_text = ocr_pdf(uploaded_pdf)
        st.subheader("📄 Extracted Text")
        st.text_area("Extracted text", value=essay_text, height=250)

st.write("---")
if st.button("🔍 Evaluate Essay"):
    if not essay_text or not essay_text.strip():
        st.error("Please provide essay text (paste/upload/camera/pdf) before evaluating.")
    else:
        with st.spinner("Analyzing with AI..."):
            result = ai_evaluate(essay_text, score_scale=10)

        # Show criteria scores
        st.subheader("📊 Dynamic Criteria & Scores (1–10)")
        cols = st.columns(2)
        # Display as list with explanation
        for c in result.get("criteria", []):
            cols[0].markdown(f"**{c['name']}** — **{c['score']}/10**")
            cols[1].markdown(f"{c.get('explanation','')}")

        st.write("---")
        # Corrected Essay (plain)
        st.subheader("✔ Corrected Essay")
        corrected = result.get("corrected_essay", essay_text)
        st.text_area("Corrected essay (no highlights)", value=corrected, height=300)

        # Summary
        st.subheader("📑 Summary Analysis")
        st.text_area("Summary", value=result.get("summary", ""), height=120)

        # Explanations (expandable)
        st.subheader("📘 Teaching Mode — Sentence-by-sentence explanations")
        st.expander("Show explanations").write(result.get("explanations", ""))

        # Download options
        st.write("---")
        st.subheader("📥 Download report / corrected essay")
        # Plain text:
        st.download_button("Download corrected essay (.txt)",
                           data=corrected,
                           file_name="corrected_essay.txt",
                           mime="text/plain")
        # Word doc:
        try:
            doc = Document()
            doc.add_heading("AI Essay Evaluation Report", level=1)
            doc.add_heading("Corrected Essay", level=2)
            doc.add_paragraph(corrected)
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(result.get("summary",""))
            doc.add_heading("Criteria & Scores", level=2)
            for c in result.get("criteria", []):
                doc.add_paragraph(f"{c['name']} — {c['score']}/10\n{c.get('explanation','')}")
            doc.add_heading("Explanations", level=2)
            doc.add_paragraph(result.get("explanations",""))
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            st.download_button("Download report (.docx)", data=doc_bytes, file_name="essay_report.docx")
        except Exception as e:
            st.warning("Docx download not available: " + str(e))

        # JSON full report
        st.download_button("Download full JSON report", data=json.dumps(result, indent=2), file_name="essay_report.json", mime="application/json")

        st.success("Thank you for using the AI Essay Checker!")
